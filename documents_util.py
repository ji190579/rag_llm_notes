import os
import re
import io
from docx import Document  # For handling .docx files
import uuid
import pdfplumber
import docx
import pytesseract
import cv2

# import fitz
# from PIL import Image
#from langchain.document_loaders import TextLoader

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter

from langchain_openai import ChatOpenAI

import nltk
from collections import Counter
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import string
import time

llm = ChatOpenAI(
    model="meta-llama/Llama-3.3-70B-Instruct-Turbo-Free",
    openai_api_key="",
    openai_api_base="https://api.together.xyz/v1",
    temperature=0.5
)



pytesseract.pytesseract.tesseract_cmd = r"C:\Users\lb02304\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"

import os
import fitz
import pytesseract
from PIL import Image
import pandas as pd
import io
import json
from tqdm import tqdm

def is_table_like(df):
    if df.empty or len(df['line_num'].unique()) < 2:
        return False
    word_counts = df.groupby("line_num")["text"].count().values
    return all(w >= 2 for w in word_counts[:3])

def extract_text_and_images_with_ocr(pdf_path, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    doc = fitz.open(pdf_path)
    all_pages = []

    #for page_index, page in enumerate(doc):
    #for page_index, page in enumerate(tqdm(doc, desc="📄 Processing pages")):
    progress = tqdm(doc, desc="📄 Processing pages")
    for page_index, page in enumerate(progress):
        progress.set_postfix_str(f"Page {page_index + 1}")  
        blocks = page.get_text("dict")["blocks"]
        page_data = []
        used_images = set()

        for block in blocks:
            bbox = block["bbox"]

            if block["type"] == 0:
                text = ""
                for line in block["lines"]:
                    for span in line["spans"]:
                        text += span["text"] + " "
                page_data.append({
                    "type": "text",
                    "content": text.strip(),
                    "bbox": bbox
                })

            elif block["type"] == 1:
                images = page.get_images(full=True)
                for img_index, img in enumerate(images):
                    xref = img[0]
                    if xref in used_images:
                        continue
                    used_images.add(xref)

                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    image = Image.open(io.BytesIO(image_bytes))

                    image_filename = f"page_{page_index + 1}_img_{img_index + 1}.{image_ext}"
                    image_path = os.path.join(output_folder, image_filename)
                    image.save(image_path)

                    ocr_text = pytesseract.image_to_string(image)
                    df = pytesseract.image_to_data(image, output_type=pytesseract.Output.DATAFRAME)
                    df = df.dropna().query("text != ''")

                    table_rows = []
                    type_hint = "unknown"

                    if is_table_like(df):
                        type_hint = "table"
                        for line_num in sorted(df['line_num'].unique()):
                            line_words = df[df['line_num'] == line_num]['text'].tolist()
                            table_rows.append(" ".join(line_words))
                    elif "→" in ocr_text or "→" in ''.join(df["text"].values):
                        type_hint = "flowchart"
                    elif any(k in ocr_text.lower() for k in ["axis", "graph", "curve", "slope", "x=", "y="]):
                        type_hint = "math_graph"
                    elif len(ocr_text.split()) < 5:
                        type_hint = "label_or_icon"
                    else:
                        type_hint = "text_image"

                    page_data.append({
                        "type": "image",
                        "type_hint": type_hint,
                        "ocr_text": ocr_text.strip(),
                        "ocr_table": table_rows if type_hint == "table" else None,
                        "bbox": bbox
                    })
                    break

        page_data_sorted = sorted(page_data, key=lambda x: x["bbox"][1])
        all_pages.append({
            "page": page_index + 1,
            "items": page_data_sorted
        })

    return all_pages

def merge_text_blocks(all_pages, y_threshold=50):
    merged = []

    for page in all_pages:
        text_blocks = [b for b in page["items"] if b["type"] == "text"]
        text_blocks_sorted = sorted(text_blocks, key=lambda x: x["bbox"][1])

        current = []
        for block in text_blocks_sorted:
            if not current:
                current.append(block)
            else:
                prev = current[-1]
                if abs(block["bbox"][1] - prev["bbox"][3]) < y_threshold:
                    current.append(block)
                else:
                    merged.append(" ".join(b["content"] for b in current))
                    current = [block]

        if current:
            merged.append(" ".join(b["content"] for b in current))

    return merged
def extract_text_from_pdf_v2(pdf_path):
    output_folder = "temp_images"
    pages_data = extract_text_and_images_with_ocr(pdf_path, output_folder)
    merged_text_blocks = merge_text_blocks(pages_data)
    return "\n\n".join(merged_text_blocks).strip()
# def extract_text_from_pdf_with_ocr(pdf_path):
#     doc = fitz.open(pdf_path)
#     combined_text = ""

#     for page_num in range(len(doc)):
#         page = doc.load_page(page_num)
#         text = page.get_text()
#         combined_text += text + "\n"

#         images = page.get_images(full=True)
#         for img_index, img in enumerate(images):
#             xref = img[0]
#             base_image = doc.extract_image(xref)
#             image_bytes = base_image["image"]
#             img_ext = base_image["ext"]

#             image = Image.open(io.BytesIO(image_bytes))
#             ocr_text = pytesseract.image_to_string(image)
#             combined_text += "\n[OCR IMAGE TEXT]\n" + ocr_text + "\n"

#     return combined_text

def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text.strip()

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="utf-8") as file:
        return file.read()

def extract_code_from_image(image_path):
    """ Extract text from code screenshots using OCR """
    image = cv2.imread(image_path)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)  # Convert to grayscale
    text = pytesseract.image_to_string(gray)  # Extract text
    return text
def get_first_line_title(text, file_name):
    if not text:
        return os.path.splitext(file_name)[0]  # fallback title

    lines = text.splitlines()
    for line in lines:
        clean_line = line.strip()
        if clean_line:  # skip empty lines
            return clean_line

    return os.path.splitext(file_name)[0]  # fallback again if all lines are empty
# Load all documents from a folder
import os
from tqdm import tqdm
from langchain.schema import Document  # ✅ Make sure this is correct


def get_first_line_title(text, file_name, max_length=120):
    if not text:
        return os.path.splitext(file_name)[0][:max_length]

    lines = text.splitlines()
    for line in lines:
        clean_line = line.strip()
        if clean_line and len(clean_line) <= max_length:
            return clean_line

    return os.path.splitext(file_name)[0][:max_length]

def load_documents(folder_path):
    documents = []

    for file in tqdm(os.listdir(folder_path), desc="📂 Loading files"):
        file_path = os.path.join(folder_path, file)

        # 👇 Extract text first — don't append here
        if file.endswith(".pdf"):
            text = extract_text_from_pdf_v2(file_path)

        elif file.endswith(".docx"):
            text = extract_text_from_docx(file_path)

        elif file.endswith(".txt"):
            text = extract_text_from_txt(file_path)

        elif file.endswith((".png", ".jpg", ".jpeg")):
            text = extract_code_from_image(file_path)

        else:
            continue  # Unsupported type

        if not text:
            print(f"Skipping empty or unreadable file: {file}")
            continue

        # 👇 Build title and wrap in Document
        title = get_first_line_title(text, file)
        doc = Document(page_content=text, metadata={"source": file, "title": title})
        documents.append(doc)

    return documents

def load_documents_v2(folder_path):
    documents = []
    prev_chapter = None  # Initialize prev_chapter before the loop
    
    for file in tqdm(os.listdir(folder_path), desc="📂 Loading files"):
        file_path = os.path.join(folder_path, file)

        # Extract full text from each file
        if file.endswith(".pdf"):
            content = extract_text_from_pdf_v2(file_path)
        elif file.endswith(".docx"):
            content = extract_text_from_docx(file_path)
        elif file.endswith(".txt"):
            content = extract_text_from_txt(file_path)
        elif file.endswith((".png", ".jpg", ".jpeg")):
            content = extract_code_from_image(file_path)
        else:
            continue

        # 📌 Detect chapter title or use fallback to previous chapter
        chapter_title = extract_chapter_from_page(content)
        
        # If no chapter found and we have a previous chapter, use it
        if not chapter_title and prev_chapter:
            chapter_title = prev_chapter
        
        if content:
            documents.append(Document(
                page_content=content,
                metadata={
                    "source": file,
                    "chapter_title": chapter_title if chapter_title else "Unknown Chapter"
                }
            ))
        
        # Update prev_chapter only if we found a new chapter
        if chapter_title:
            prev_chapter = chapter_title
    
    return documents




# fkind first chapter /couser/lesson title to the chunke
import re


def find_first_chapter_title_v2(text, mode="chapter", max_line_length=120):
    lines = text.splitlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 🔍 Step 1: Try known chapter patterns
        if mode in ("chapter", "flexible"):
            patterns = [
                r"(?i)\bchapter\s*\d+[:\-\.\s]*(.+)",
                r"(?i)\b\d+\.\s+(.+)",
                r"(?i)\bcourse[:\-\.\s]*(.+)",
                r"(?i)\blesson\s*\d+[:\-\.\s]*(.+)",
                r"(?i)^(.+):\s*$"  # Line ends in colon with a title
            ]

            for pattern in patterns:
                match = re.search(pattern, line)
                if match:
                    return match.group(1).strip()

        # 🧪 Step 2: Use first non-empty line (with or without `:`)
        if mode in ("first_line", "flexible"):
            if len(line) <= max_line_length:
                return line.rstrip(":").strip()

    return None





def extract_chapter_from_page(page_content: str) -> str:
    """
    Extract chapter title from a single page/document.
    
    Looks for patterns:
    1. "Chapter:" + separator + title (e.g., "Chapter: Microservices", "Chapter 1: Microservices")
    2. Number + separator + title (e.g., "1. Microservices", "2 - Implementation")
    
    Args:
        page_content (str): Content of a single page/document
        
    Returns:
        str: Chapter title if found, empty string if not found
    """
    if not page_content or not page_content.strip():
        return ""
    
    lines = page_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Pattern 1: "Chapter" + optional number + separator + title
        # Matches: "Chapter: Microservices", "Chapter 1: Microservices", "Chapter - Implementation"
        chapter_pattern = r'^\s*chapter\s*(?:\d+)?\s*[\:\-\.\s]+(.+?)$'
        match = re.match(chapter_pattern, line, re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            if title:
                return title
        
        # Pattern 2: Whole number + separator + title (NOT decimal numbers like 1.1, 1.22)
        # Matches: "1. Microservices", "2 - Implementation", "3: Use Cases", "1 Microservices"
        # But NOT: "1.1 What are...", "1.22 Details..."
        number_pattern = r'^\s*(\d+)\s*[\:\-\s]+(.+?)$|^\s*(\d+)\.(?!\d)\s*(.+?)$'
        match = re.match(number_pattern, line, re.IGNORECASE)
        if match:
            # Handle both pattern groups (with different separators)
            title = (match.group(2) or match.group(4)).strip() if match.group(2) or match.group(4) else ""
            if title:
                return title
    
    return ""



nltk.download("punkt_tab")  # Changed from "punkt" to "punkt_tab"
nltk.download("stopwords")

def extract_keywords(text, top_n=5):
    stop_words = set(stopwords.words("english"))
    words = word_tokenize(text.lower())
    words = [word for word in words if word.isalnum() and word not in stop_words]
    
    most_common = Counter(words).most_common(top_n)
    keywords = [word for word, _ in most_common]
    return keywords



# here i used manual semantic chunking by: page or logic  logic separator,summary from prev chunk,keyword
from langchain.text_splitter import RecursiveCharacterTextSplitter
import uuid
import time

def chunk_documents(docs, chunk_size=500, chunk_overlap=50, page=False, summary_chunk=False, keywords_chunk=False):
    chunked_docs = []
    ## if page is as true the give prioity for page then logical separators
     ## if page is as false the give the logical separators

    separators = ["page", "\n\n", "\n", ".", " ", ""] if page else ["\n\n", "\n", ".", " ", ""]

    text_splitter = RecursiveCharacterTextSplitter(
        separators=separators,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )

    if isinstance(docs, str):
        docs = [Document(page_content=docs, metadata={})]  # Wrap raw text in a Document

    if isinstance(docs, list) and all(isinstance(doc, Document) for doc in docs):
        chunks = text_splitter.split_documents(docs)
    else:
        docs = [Document(page_content=str(doc), metadata={}) for doc in docs]
        chunks = text_splitter.split_documents(docs)

    prev_summary = None
    prev_keywords = None

    for idx, chunk in enumerate(chunks):
        if not chunk.page_content.strip():
            continue

        summary = (
            llm.invoke(f"Summarize this text in 1 sentence:\n\n{chunk.page_content}").content
            if summary_chunk else ""
        )

        keywords = extract_keywords(chunk.page_content) if keywords_chunk else ""

        chapter_title = find_first_chapter_title_v2(chunk.page_content)
        
        #  If found, update the current chapter
        if chapter_title:
            current_chapter = chapter_title

        # Get original source and title from the metadata
        source = chunk.metadata.get("source", "unknown_source")
        title = chunk.metadata.get("title", "unknown_title")
        chapter_title=title
        chunked_docs.append({
            "id": str(uuid.uuid4()),
            "chapter": current_chapter,
            "source": source,
            "title": title,
            "text": chunk.page_content,
            "metadata": {
                "summary": summary,
                "prev_summary": prev_summary,
                "chunk_index": idx,
                "chunk_length": len(chunk.page_content),
                "prev_keywords": prev_keywords,
                "keywords": keywords
            }
        })

        prev_summary = summary
        prev_keywords = keywords

    return chunked_docs

        

import json
import pprint

def print_data_as_format_json(docs):
    pp = pprint.PrettyPrinter(
        indent=2, 
        width=120,        # Wider output
        depth=None,       # No depth limit
        compact=False     # Don't compact
    )
    pp.pprint(docs)
